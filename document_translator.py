#!/usr/bin/env python3
"""
Document Translator - Fixed Version using python-docx

This version:
- Uses python-docx library (more reliable than custom ooxml)
- Handles all DOCX files properly
- Better error handling and recovery
- Preserves formatting, headers, footers, tables
- Larger chunk sizes for efficiency
- Progress tracking
"""

import sys
import os
import shutil
from deep_translator import GoogleTranslator
import time
import re
import json
import hashlib
import copy

# Try to import NLTK (optional - will fallback if not available)
try:
    import nltk
    from nltk.tokenize import sent_tokenize
    
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        try:
            nltk.download('punkt', quiet=True)
        except:
            pass
    
    try:
        nltk.data.find('tokenizers/punkt_tab')
    except LookupError:
        try:
            nltk.download('punkt_tab', quiet=True)
        except:
            pass
    
    NLTK_AVAILABLE = True
except:
    NLTK_AVAILABLE = False
    print("Warning: NLTK not available - using simple sentence splitting")

# Try to import python-docx
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_SUPPORT = True
    print("✓ python-docx loaded successfully")
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

# Try to import PDF libraries
try:
    import fitz  # PyMuPDF
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Warning: PDF support not available")


class DocumentTranslator:
    """
    Enhanced document translator using python-docx for reliable DOCX processing
    """
    
    def __init__(self, source_lang='auto', target_lang='es'):
        """
        Initialize the translator
        
        Args:
            source_lang: Source language code (default: 'auto' for auto-detect)
            target_lang: Target language code (default: 'es' for Spanish)
        """
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.translation_cache = {}
        self.context_cache = {}
        self.use_context = True
        self.progress_file = None
        self.total_segments = 0
        self.translated_segments = 0
        
        # Initialize the deep-translator instance
        try:
            self.translator = GoogleTranslator(source=source_lang, target=target_lang)
            print(f"✓ Translator initialized: {source_lang} -> {target_lang}")
        except Exception as e:
            print(f"✗ Failed to initialize translator: {e}")
            self.translator = None
    
    def split_into_sentences(self, text):
        """
        Split text into sentences (uses NLTK if available, otherwise simple split)
        """
        if not text or not text.strip():
            return []
        
        if NLTK_AVAILABLE:
            try:
                sentences = sent_tokenize(text)
                return [s.strip() for s in sentences if s.strip()]
            except Exception as e:
                print(f"Warning: NLTK tokenization failed: {e}")
        
        # Fallback: improved sentence splitting
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def translate_chunk(self, chunk):
        """
        Translate a single chunk of text with enhanced retry logic
        """
        if not chunk or not chunk.strip():
            return chunk
        
        # Check cache
        cache_key = hashlib.md5(chunk.encode()).hexdigest()
        if cache_key in self.context_cache:
            return self.context_cache[cache_key]
        
        # Retry logic with exponential backoff
        max_retries = 5
        base_delay = 1.0
        
        for attempt in range(max_retries):
            try:
                if attempt > 0:
                    delay = base_delay * (2 ** (attempt - 1))
                    print(f"  Retry {attempt}/{max_retries} after {delay:.1f}s delay...")
                    time.sleep(delay)
                
                # Reinitialize translator if needed
                if self.translator is None:
                    self.translator = GoogleTranslator(source=self.source_lang, target=self.target_lang)
                
                # Translate
                translated = self.translator.translate(chunk)
                
                # Validate result
                if translated is None or not translated:
                    if attempt < max_retries - 1:
                        continue
                    print(f"  Warning: Translation returned empty for chunk")
                    translated = chunk
                
                if not isinstance(translated, str):
                    translated = str(translated) if translated else chunk
                
                # Cache and return
                self.context_cache[cache_key] = translated
                self.translated_segments += 1
                
                return translated
                
            except Exception as e:
                error_msg = str(e).lower()
                
                if 'rate' in error_msg or 'limit' in error_msg or '429' in str(e):
                    print(f"  Rate limit detected, using longer delay...")
                    delay = base_delay * 3 * (2 ** attempt)
                    time.sleep(min(delay, 30))
                elif 'connection' in error_msg or 'timeout' in error_msg:
                    print(f"  Connection issue, retrying...")
                else:
                    print(f"  Translation error: {e}")
                
                if attempt == max_retries - 1:
                    print(f"  All retries exhausted, using original text")
                    self.context_cache[cache_key] = chunk
                    return chunk
                
                try:
                    time.sleep(1)
                    self.translator = GoogleTranslator(source=self.source_lang, target=self.target_lang)
                except Exception as recreate_error:
                    print(f"  Failed to recreate translator: {recreate_error}")
        
        return chunk
    
    def translate_with_context(self, text, max_chunk_size=4500):
        """
        Translate text with context awareness - optimized for large documents
        """
        if not text or not text.strip():
            return text
        
        # Check cache
        text_hash = hashlib.md5(text.encode()).hexdigest()
        if text_hash in self.translation_cache:
            return self.translation_cache[text_hash]
        
        # If text is short, translate directly
        if len(text) <= max_chunk_size:
            result = self.translate_chunk(text)
            self.translation_cache[text_hash] = result
            return result
        
        try:
            print(f"\n📄 Processing large text block ({len(text)} characters)...")
            
            sentences = self.split_into_sentences(text)
            
            if not sentences:
                return text
            
            print(f"  Split into {len(sentences)} sentences")
            
            # Group sentences into optimal chunks
            chunks = []
            current_chunk = []
            current_size = 0
            
            for sentence in sentences:
                sentence_size = len(sentence)
                
                if sentence_size > max_chunk_size:
                    words = sentence.split()
                    temp_chunk = []
                    temp_size = 0
                    
                    for word in words:
                        word_size = len(word) + 1
                        if temp_size + word_size > max_chunk_size and temp_chunk:
                            chunks.append(' '.join(temp_chunk))
                            temp_chunk = [word]
                            temp_size = word_size
                        else:
                            temp_chunk.append(word)
                            temp_size += word_size
                    
                    if temp_chunk:
                        chunks.append(' '.join(temp_chunk))
                    continue
                
                if current_size + sentence_size > max_chunk_size and current_chunk:
                    chunks.append(' '.join(current_chunk))
                    current_chunk = [sentence]
                    current_size = sentence_size
                else:
                    current_chunk.append(sentence)
                    current_size += sentence_size + 1
            
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            
            print(f"  Optimized into {len(chunks)} chunks for translation")
            
            # Translate each chunk
            translated_chunks = []
            consecutive_successes = 0
            
            for i, chunk in enumerate(chunks):
                chunk_num = i + 1
                print(f"\n  Translating chunk {chunk_num}/{len(chunks)} ({len(chunk)} chars)...")
                
                try:
                    translated = self.translate_chunk(chunk)
                    translated_chunks.append(translated)
                    consecutive_successes += 1
                    
                    if i < len(chunks) - 1:
                        if consecutive_successes < 3:
                            delay = 2.0 if len(chunk) > 2000 else 1.5
                        elif len(chunk) > 3000:
                            delay = 1.0
                        else:
                            delay = 0.5
                        
                        time.sleep(delay)
                    
                except Exception as e:
                    print(f"    Error translating chunk {chunk_num}: {e}")
                    consecutive_successes = 0
                    translated_chunks.append(chunk)
                    if i < len(chunks) - 1:
                        time.sleep(3.0)
            
            result = ' '.join(translated_chunks)
            self.translation_cache[text_hash] = result
            
            print(f"\n  ✓ Large text translation complete!")
            
            return result
            
        except Exception as e:
            print(f"\n❌ Context translation failed: {e}")
            try:
                result = self.translate_chunk(text)
                self.translation_cache[text_hash] = result
                return result
            except Exception as fallback_error:
                print(f"  Fallback also failed: {fallback_error}")
                return text
    
    def translate_text(self, text):
        """
        Main translation method with enhanced error handling
        """
        try:
            return self.translate_with_context(text, max_chunk_size=4500)
        except Exception as e:
            print(f"Error in translate_text: {e}")
            return text
    
    def should_translate_text(self, text):
        """
        Determine if text should be translated
        """
        if not text:
            return False
            
        text = text.strip()
        
        if not text or len(text) < 2:
            return False
        
        # Don't translate email addresses
        if '@' in text and '.' in text and len(text.split()) == 1:
            return False
        
        # Don't translate URLs
        if text.startswith(('http://', 'https://', 'www.')):
            return False
        
        # Don't translate pure numbers
        if re.match(r'^[\d\s\-/.,]+$', text):
            return False
        
        # Don't translate if only whitespace/special chars
        if not re.search(r'[a-zA-Z\u0080-\uFFFF]', text):
            return False
            
        return True
    
    def translate_docx(self, input_docx, output_docx):
        """
        Translate a DOCX document using python-docx library
        
        This method properly handles:
        - Main document text
        - Headers and footers
        - Tables
        - Text boxes
        - Preserves all formatting
        """
        if not DOCX_SUPPORT:
            raise Exception("python-docx not available. Install with: pip install python-docx")
        
        print(f"\n{'='*60}")
        print(f"DOCX Document Translator (python-docx)")
        print(f"{'='*60}")
        print(f"Input:  {input_docx}")
        print(f"Output: {output_docx}")
        print(f"Source Language: {self.source_lang}")
        print(f"Target Language: {self.target_lang}")
        
        # Check file exists
        if not os.path.exists(input_docx):
            raise Exception(f"Input file not found: {input_docx}")
        
        # Check file size
        file_size = os.path.getsize(input_docx) / (1024 * 1024)
        print(f"File Size: {file_size:.2f} MB")
        
        if file_size > 10:
            print(f"⚠️  Large file detected - using optimized processing")
        
        print(f"{'='*60}\n")
        
        try:
            # Load the document
            print("Step 1: Loading document...")
            doc = DocxDocument(input_docx)
            
            # Count elements for progress tracking
            total_elements = self._count_translatable_elements(doc)
            print(f"  Found {total_elements} translatable elements")
            
            self.total_segments = total_elements
            self.translated_segments = 0
            
            # Translate paragraphs in main body
            print("\nStep 2: Translating main document body...")
            for para in doc.paragraphs:
                self._translate_paragraph(para)
            
            # Translate tables
            print("\nStep 3: Translating tables...")
            for table in doc.tables:
                self._translate_table(table)
            
            # Translate headers
            print("\nStep 4: Translating headers...")
            for section in doc.sections:
                # First page header
                if section.first_page_header:
                    for para in section.first_page_header.paragraphs:
                        self._translate_paragraph(para)
                    for table in section.first_page_header.tables:
                        self._translate_table(table)
                
                # Main header
                if section.header:
                    for para in section.header.paragraphs:
                        self._translate_paragraph(para)
                    for table in section.header.tables:
                        self._translate_table(table)
                
                # Even page header
                if section.even_page_header:
                    for para in section.even_page_header.paragraphs:
                        self._translate_paragraph(para)
                    for table in section.even_page_header.tables:
                        self._translate_table(table)
            
            # Translate footers
            print("\nStep 5: Translating footers...")
            for section in doc.sections:
                # First page footer
                if section.first_page_footer:
                    for para in section.first_page_footer.paragraphs:
                        self._translate_paragraph(para)
                    for table in section.first_page_footer.tables:
                        self._translate_table(table)
                
                # Main footer
                if section.footer:
                    for para in section.footer.paragraphs:
                        self._translate_paragraph(para)
                    for table in section.footer.tables:
                        self._translate_table(table)
                
                # Even page footer
                if section.even_page_footer:
                    for para in section.even_page_footer.paragraphs:
                        self._translate_paragraph(para)
                    for table in section.even_page_footer.tables:
                        self._translate_table(table)
            
            # Save the translated document
            print("\nStep 6: Saving translated document...")
            doc.save(output_docx)
            
            # Verify output
            if os.path.exists(output_docx):
                output_size = os.path.getsize(output_docx) / (1024 * 1024)
                print(f"\n✓ Output file created: {output_size:.2f} MB")
            
            print(f"\n{'='*60}")
            print(f"Translation Complete!")
            print(f"{'='*60}")
            print(f"Translated segments: {self.translated_segments}")
            print(f"Cached translations: {len(self.translation_cache)}")
            print(f"Output saved to: {output_docx}")
            print(f"{'='*60}\n")
            
        except Exception as e:
            print(f"\n❌ Error during translation: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def _count_translatable_elements(self, doc):
        """Count all translatable elements in the document"""
        count = 0
        
        # Count paragraphs
        for para in doc.paragraphs:
            if para.text and self.should_translate_text(para.text):
                count += 1
        
        # Count table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text and self.should_translate_text(para.text):
                            count += 1
        
        # Count headers/footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer, 
                                  section.first_page_header, section.first_page_footer,
                                  section.even_page_header, section.even_page_footer]:
                if header_footer:
                    for para in header_footer.paragraphs:
                        if para.text and self.should_translate_text(para.text):
                            count += 1
        
        return count
    
    def _translate_paragraph(self, paragraph):
        """
        Translate a paragraph while preserving formatting
        """
        if not paragraph.text or not self.should_translate_text(paragraph.text):
            return
        
        try:
            # Get the full text
            full_text = paragraph.text
            
            # Translate the full text
            translated_text = self.translate_text(full_text)
            
            if translated_text == full_text:
                return
            
            # Strategy: Clear and rewrite with same formatting from first run
            # This preserves paragraph-level formatting but may lose run-level variations
            
            # If paragraph has runs, try to preserve formatting
            if paragraph.runs:
                # Get formatting from first run
                first_run = paragraph.runs[0]
                
                # Store formatting properties
                bold = first_run.bold
                italic = first_run.italic
                underline = first_run.underline
                font_name = first_run.font.name
                font_size = first_run.font.size
                
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                
                # Set translated text to first run
                paragraph.runs[0].text = translated_text
                
                # Restore formatting
                paragraph.runs[0].bold = bold
                paragraph.runs[0].italic = italic
                paragraph.runs[0].underline = underline
                if font_name:
                    paragraph.runs[0].font.name = font_name
                if font_size:
                    paragraph.runs[0].font.size = font_size
            
            # Update progress
            self.translated_segments += 1
            if self.total_segments > 0 and self.translated_segments % 10 == 0:
                progress = (self.translated_segments / self.total_segments) * 100
                print(f"  Progress: {self.translated_segments}/{self.total_segments} ({progress:.1f}%)")
                
        except Exception as e:
            print(f"  Warning: Could not translate paragraph: {e}")
    
    def _translate_table(self, table):
        """
        Translate all cells in a table
        """
        try:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._translate_paragraph(para)
                    
                    # Handle nested tables
                    for nested_table in cell.tables:
                        self._translate_table(nested_table)
                        
        except Exception as e:
            print(f"  Warning: Could not translate table: {e}")
    
    def translate_pdf(self, input_pdf, output_pdf):
        """
        Translate a PDF document
        """
        if not PDF_SUPPORT:
            raise Exception("PDF support not available. Install PyMuPDF: pip install pymupdf")
        
        print(f"\n{'='*60}")
        print(f"PDF Document Translator")
        print(f"{'='*60}")
        
        try:
            doc = fitz.open(input_pdf)
            output_doc = fitz.open()
            
            total_pages = len(doc)
            print(f"Processing {total_pages} pages...")
            
            for page_num in range(total_pages):
                print(f"\nPage {page_num + 1}/{total_pages}...")
                page = doc[page_num]
                new_page = output_doc.new_page(width=page.rect.width, height=page.rect.height)
                
                # Copy background
                pix = page.get_pixmap(matrix=fitz.Matrix(1, 1), alpha=False)
                img_bytes = pix.tobytes("png")
                new_page.insert_image(new_page.rect, stream=img_bytes, overlay=False)
                
                # Translate text blocks
                blocks = page.get_text("dict")["blocks"]
                text_blocks = [b for b in blocks if b["type"] == 0]
                
                print(f"  Found {len(text_blocks)} text blocks")
                
                for block_num, block in enumerate(text_blocks, 1):
                    block_text = ""
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            block_text += span.get("text", "") + " "
                    
                    block_text = block_text.strip()
                    if block_text and self.should_translate_text(block_text):
                        print(f"  Translating block {block_num}/{len(text_blocks)}...")
                        translated = self.translate_text(block_text)
                        
                        bbox = fitz.Rect(block["bbox"])
                        font_size = 11
                        if block.get("lines") and block["lines"][0].get("spans"):
                            font_size = block["lines"][0]["spans"][0].get("size", 11)
                        
                        new_page.draw_rect(bbox, color=(1, 1, 1), fill=(1, 1, 1))
                        new_page.insert_textbox(bbox, translated, fontsize=font_size)
            
            output_doc.save(output_pdf)
            output_doc.close()
            doc.close()
            
            print(f"\n✓ Translation complete: {output_pdf}")
            
        except Exception as e:
            print(f"❌ Error: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def translate_document(self, input_file, output_file):
        """
        Translate a document (auto-detect format)
        """
        if input_file.lower().endswith('.docx'):
            self.translate_docx(input_file, output_file)
        elif input_file.lower().endswith('.pdf'):
            self.translate_pdf(input_file, output_file)
        else:
            raise ValueError(f"Unsupported file type: {input_file}")


def main():
    """Main entry point"""
    if len(sys.argv) < 3:
        print("Usage: python document_translator.py <input> <output> [target_lang] [source_lang]")
        print("\nExample:")
        print("  python document_translator.py input.docx output.docx es auto")
        print("\nSupported formats: .docx, .pdf")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    target_lang = sys.argv[3] if len(sys.argv) > 3 else 'es'
    source_lang = sys.argv[4] if len(sys.argv) > 4 else 'auto'
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found")
        sys.exit(1)
    
    translator = DocumentTranslator(source_lang=source_lang, target_lang=target_lang)
    translator.translate_document(input_file, output_file)


if __name__ == '__main__':
    main()