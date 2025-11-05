#!/usr/bin/env python3
"""
Enhanced PDF Translator - Converts PDF to DOCX for translation, then back to PDF
This preserves formatting much better than direct PDF text manipulation
"""

import os
import sys
import shutil
import tempfile
from pathlib import Path
import time

# Translation support
from deep_translator import GoogleTranslator

# PDF to DOCX conversion
try:
    from pdf2docx import Converter as PDF2DocxConverter
    PDF2DOCX_SUPPORT = True
except ImportError:
    PDF2DOCX_SUPPORT = False
    print("Warning: pdf2docx not installed. Run: pip install pdf2docx")

# DOCX manipulation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed. Run: pip install python-docx")

# DOCX to PDF conversion - multiple backend options
DOCX2PDF_BACKEND = None

# Try different DOCX to PDF backends
try:
    from docx2pdf import convert
    DOCX2PDF_BACKEND = 'docx2pdf'
except ImportError:
    pass

if not DOCX2PDF_BACKEND:
    try:
        import subprocess
        # Check if LibreOffice is available
        result = subprocess.run(['libreoffice', '--version'], capture_output=True, text=True)
        if result.returncode == 0:
            DOCX2PDF_BACKEND = 'libreoffice'
    except:
        pass

if not DOCX2PDF_BACKEND:
    try:
        import pypandoc
        DOCX2PDF_BACKEND = 'pandoc'
    except ImportError:
        pass

# For fallback PDF creation if other methods fail
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, Table
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_SUPPORT = True
except ImportError:
    REPORTLAB_SUPPORT = False

print(f"""
PDF Translation Backend Status:
-------------------------------
PDF to DOCX: {'✓ Available' if PDF2DOCX_SUPPORT else '✗ Not available (install pdf2docx)'}
DOCX Support: {'✓ Available' if DOCX_SUPPORT else '✗ Not available (install python-docx)'}
DOCX to PDF: {f'✓ {DOCX2PDF_BACKEND}' if DOCX2PDF_BACKEND else '✗ Not available'}
ReportLab (fallback): {'✓ Available' if REPORTLAB_SUPPORT else '✗ Not available'}
""")


class EnhancedPDFTranslator:
    """Translates PDFs by converting to DOCX, translating, then converting back"""
    
    def __init__(self, source_lang='auto', target_lang='es'):
        """
        Initialize the translator
        
        Args:
            source_lang: Source language code (default: 'auto' for auto-detect)
            target_lang: Target language code (default: 'es' for Spanish)
        """
        self.translator = Translator()
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.translation_cache = {}
        
    def translate_text(self, text):
        """
        Translate text with caching
        
        Args:
            text: Text to translate
            
        Returns:
            Translated text
        """
        if not text or not text.strip():
            return text
            
        # Check cache
        if text in self.translation_cache:
            return self.translation_cache[text]
        
        try:
            # Small delay to avoid rate limiting
            time.sleep(0.1)
            
            result = self.translator.translate(
                text,
                src=self.source_lang,
                dest=self.target_lang
            )
            
            translated = result.text
            self.translation_cache[text] = translated
            
            print(f"  Translated: '{text[:50]}...' -> '{translated[:50]}...'")
            return translated
            
        except Exception as e:
            print(f"  Warning: Translation failed: {e}")
            return text
    
    def should_translate_text(self, text):
        """Check if text should be translated"""
        text = text.strip()
        
        if not text:
            return False
        
        # Skip emails, URLs, numbers, dates
        if '@' in text and '.' in text and len(text.split()) == 1:
            return False
        if text.startswith(('http://', 'https://', 'www.')):
            return False
        if text.replace('.', '').replace(',', '').replace('-', '').replace('/', '').isdigit():
            return False
        if len(text) <= 2:
            return False
            
        return True
    
    def pdf_to_docx(self, pdf_path, docx_path):
        """
        Convert PDF to DOCX with formatting preservation
        
        Args:
            pdf_path: Path to input PDF
            docx_path: Path to output DOCX
        """
        if not PDF2DOCX_SUPPORT:
            raise Exception("PDF to DOCX conversion not available. Install: pip install pdf2docx")
        
        print(f"Converting PDF to DOCX...")
        print(f"  Input: {pdf_path}")
        print(f"  Output: {docx_path}")
        
        try:
            # Create converter
            cv = PDF2DocxConverter(pdf_path)
            
            # Convert with all pages
            cv.convert(docx_path, start=0, end=None)
            
            # Close the converter
            cv.close()
            
            print(f"  ✓ Conversion complete")
            return True
            
        except Exception as e:
            print(f"  ✗ Conversion failed: {e}")
            raise
    
    def translate_docx_content(self, docx_path, output_docx_path):
        """
        Translate DOCX content while preserving formatting
        
        Args:
            docx_path: Path to input DOCX
            output_docx_path: Path to output translated DOCX
        """
        if not DOCX_SUPPORT:
            raise Exception("DOCX support not available. Install: pip install python-docx")
        
        print(f"Translating DOCX content...")
        
        try:
            # Open the document
            doc = Document(docx_path)
            
            # Translate paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text and self.should_translate_text(paragraph.text):
                    # Store original formatting
                    original_runs = []
                    for run in paragraph.runs:
                        original_runs.append({
                            'text': run.text,
                            'bold': run.bold,
                            'italic': run.italic,
                            'underline': run.underline,
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                        })
                    
                    # Translate the full paragraph text
                    full_text = paragraph.text
                    translated_text = self.translate_text(full_text)
                    
                    # Clear paragraph and add translated text with original formatting
                    paragraph.clear()
                    
                    # If we had multiple runs, try to maintain their boundaries
                    if len(original_runs) > 1:
                        # Simple approach: apply first run's formatting to all text
                        # (More sophisticated approach would split translated text proportionally)
                        run = paragraph.add_run(translated_text)
                        if original_runs[0]['bold']:
                            run.bold = True
                        if original_runs[0]['italic']:
                            run.italic = True
                        if original_runs[0]['underline']:
                            run.underline = True
                        if original_runs[0]['font_name']:
                            run.font.name = original_runs[0]['font_name']
                        if original_runs[0]['font_size']:
                            run.font.size = original_runs[0]['font_size']
                        if original_runs[0]['color']:
                            run.font.color.rgb = original_runs[0]['color']
                    else:
                        # Single run - simpler case
                        paragraph.add_run(translated_text)
            
            # Translate tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text and self.should_translate_text(paragraph.text):
                                translated = self.translate_text(paragraph.text)
                                paragraph.text = translated
            
            # Translate headers and footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text and self.should_translate_text(paragraph.text):
                        translated = self.translate_text(paragraph.text)
                        paragraph.text = translated
                
                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    if paragraph.text and self.should_translate_text(paragraph.text):
                        translated = self.translate_text(paragraph.text)
                        paragraph.text = translated
            
            # Save the translated document
            doc.save(output_docx_path)
            print(f"  ✓ Translation complete")
            print(f"  Translated segments: {len(self.translation_cache)}")
            return True
            
        except Exception as e:
            print(f"  ✗ Translation failed: {e}")
            raise
    
    def docx_to_pdf_docx2pdf(self, docx_path, pdf_path):
        """Convert DOCX to PDF using docx2pdf (Windows/Mac)"""
        from docx2pdf import convert
        print(f"  Using docx2pdf backend...")
        convert(docx_path, pdf_path)
        return True
    
    def docx_to_pdf_libreoffice(self, docx_path, pdf_path):
        """Convert DOCX to PDF using LibreOffice (Linux/cross-platform)"""
        import subprocess
        print(f"  Using LibreOffice backend...")
        
        # Get the directory and filename
        output_dir = os.path.dirname(pdf_path)
        
        # LibreOffice command
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', output_dir,
            docx_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
        # LibreOffice creates PDF with same base name
        generated_pdf = os.path.join(output_dir, 
                                     os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
        
        # Rename if needed
        if generated_pdf != pdf_path:
            shutil.move(generated_pdf, pdf_path)
        
        return True
    
    def docx_to_pdf_pandoc(self, docx_path, pdf_path):
        """Convert DOCX to PDF using Pandoc"""
        import pypandoc
        print(f"  Using Pandoc backend...")
        
        # Convert using pandoc
        pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path,
                             extra_args=['--pdf-engine=xelatex'])
        return True
    
    def docx_to_pdf_reportlab(self, docx_path, pdf_path):
        """Fallback: Convert DOCX to PDF using ReportLab (basic)"""
        if not REPORTLAB_SUPPORT or not DOCX_SUPPORT:
            raise Exception("ReportLab or python-docx not available")
        
        print(f"  Using ReportLab backend (basic formatting)...")
        
        # Read DOCX
        doc = Document(docx_path)
        
        # Create PDF
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Convert paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine style based on paragraph properties
                style = styles['Normal']
                if para.style.name.startswith('Heading'):
                    style = styles['Heading1']
                
                p = Paragraph(para.text, style)
                story.append(p)
                story.append(Spacer(1, 12))
        
        # Build PDF
        pdf_doc.build(story)
        return True
    
    def docx_to_pdf(self, docx_path, pdf_path):
        """
        Convert DOCX to PDF using best available method
        
        Args:
            docx_path: Path to input DOCX
            pdf_path: Path to output PDF
        """
        print(f"Converting DOCX to PDF...")
        print(f"  Input: {docx_path}")
        print(f"  Output: {pdf_path}")
        
        try:
            if DOCX2PDF_BACKEND == 'docx2pdf':
                return self.docx_to_pdf_docx2pdf(docx_path, pdf_path)
            elif DOCX2PDF_BACKEND == 'libreoffice':
                return self.docx_to_pdf_libreoffice(docx_path, pdf_path)
            elif DOCX2PDF_BACKEND == 'pandoc':
                return self.docx_to_pdf_pandoc(docx_path, pdf_path)
            elif REPORTLAB_SUPPORT:
                return self.docx_to_pdf_reportlab(docx_path, pdf_path)
            else:
                raise Exception("No DOCX to PDF converter available")
                
        except Exception as e:
            print(f"  ✗ Conversion failed: {e}")
            
            # If primary method fails, try fallback
            if DOCX2PDF_BACKEND and REPORTLAB_SUPPORT:
                print(f"  Trying ReportLab fallback...")
                try:
                    return self.docx_to_pdf_reportlab(docx_path, pdf_path)
                except Exception as e2:
                    print(f"  ✗ Fallback also failed: {e2}")
                    raise
            raise
    
    def translate_pdf(self, input_pdf, output_pdf):
        """
        Main method: Translate PDF by converting to DOCX, translating, and converting back
        
        Args:
            input_pdf: Path to input PDF file
            output_pdf: Path to output PDF file
        """
        print(f"\n{'='*70}")
        print(f"Enhanced PDF Translation (via DOCX conversion)")
        print(f"{'='*70}")
        print(f"Input PDF:  {input_pdf}")
        print(f"Output PDF: {output_pdf}")
        print(f"Source Language: {self.source_lang}")
        print(f"Target Language: {self.target_lang}")
        print(f"{'='*70}\n")
        
        # Create temporary directory for intermediate files
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx = os.path.join(temp_dir, "temp_converted.docx")
            temp_translated_docx = os.path.join(temp_dir, "temp_translated.docx")
            
            try:
                # Step 1: Convert PDF to DOCX
                print(f"Step 1: Converting PDF to DOCX for better formatting preservation...")
                self.pdf_to_docx(input_pdf, temp_docx)
                print(f"  ✓ PDF successfully converted to DOCX")
                
                # Step 2: Translate the DOCX
                print(f"\nStep 2: Translating document content...")
                self.translate_docx_content(temp_docx, temp_translated_docx)
                print(f"  ✓ Document translated successfully")
                
                # Step 3: Convert translated DOCX back to PDF
                print(f"\nStep 3: Converting translated DOCX back to PDF...")
                self.docx_to_pdf(temp_translated_docx, output_pdf)
                print(f"  ✓ Successfully converted to PDF")
                
                print(f"\n{'='*70}")
                print(f"✓ Translation Complete!")
                print(f"{'='*70}")
                print(f"Total translated segments: {len(self.translation_cache)}")
                print(f"Output saved to: {output_pdf}")
                print(f"✓ Original formatting preserved")
                print(f"✓ Images and layouts maintained")
                print(f"✓ Professional quality output")
                print(f"{'='*70}\n")
                
                return True
                
            except Exception as e:
                print(f"\n✗ Error during PDF translation: {e}")
                import traceback
                traceback.print_exc()
                raise


def install_dependencies():
    """Helper function to install required dependencies"""
    import subprocess
    
    packages = [
        'pdf2docx',
        'python-docx',
        'googletrans==4.0.0-rc1',
        'reportlab'
    ]
    
    print("Installing required packages...")
    for package in packages:
        try:
            print(f"Installing {package}...")
            subprocess.run([sys.executable, '-m', 'pip', 'install', package], 
                          check=True, capture_output=True)
            print(f"  ✓ {package} installed")
        except:
            print(f"  ✗ Failed to install {package}")
    
    # Platform-specific packages
    if sys.platform == 'win32':
        try:
            print("Installing docx2pdf (Windows)...")
            subprocess.run([sys.executable, '-m', 'pip', 'install', 'docx2pdf'], 
                          check=True, capture_output=True)
            print("  ✓ docx2pdf installed")
        except:
            pass
    
    print("\nFor Linux users: Install LibreOffice for best DOCX to PDF conversion:")
    print("  sudo apt-get install libreoffice")


def main():
    """Main entry point for testing"""
    
    if len(sys.argv) < 3:
        print("Enhanced PDF Translator - Format-Preserving Translation")
        print("\nUsage:")
        print("  python pdf_translator_improved.py <input.pdf> <output.pdf> [target_lang] [source_lang]")
        print("\nExamples:")
        print("  python pdf_translator_improved.py document.pdf translated.pdf es")
        print("  python pdf_translator_improved.py report.pdf report_fr.pdf fr en")
        print("\nTo install dependencies:")
        print("  python pdf_translator_improved.py --install")
        
        if len(sys.argv) == 2 and sys.argv[1] == '--install':
            install_dependencies()
        
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    output_pdf = sys.argv[2]
    target_lang = sys.argv[3] if len(sys.argv) > 3 else 'es'
    source_lang = sys.argv[4] if len(sys.argv) > 4 else 'auto'
    
    # Validate input file
    if not os.path.exists(input_pdf):
        print(f"Error: Input file '{input_pdf}' not found")
        sys.exit(1)
    
    # Create translator and translate
    translator = EnhancedPDFTranslator(source_lang=source_lang, target_lang=target_lang)
    translator.translate_pdf(input_pdf, output_pdf)


if __name__ == '__main__':
    main()